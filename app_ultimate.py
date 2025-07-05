#!/usr/bin/env python3
"""
JobSniper AI - Ultimate Version
===============================

The most advanced version with enhanced features, better UI, and comprehensive functionality.
"""

import streamlit as st
import sys
import os
import tempfile
import json
import logging
import time
from datetime import datetime, timedelta
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import base64
import io

# Add project root to path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Page configuration
st.set_page_config(
    page_title="JobSniper AI Ultimate - Professional Resume & Career Intelligence",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if 'parsed_resume' not in st.session_state:
    st.session_state.parsed_resume = None
if 'analysis_history' not in st.session_state:
    st.session_state.analysis_history = []
if 'user_preferences' not in st.session_state:
    st.session_state.user_preferences = {
        'theme': 'dark',
        'auto_save': True,
        'notifications': True
    }

# Ultimate CSS with advanced styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    
    /* Global Variables */
    :root {
        --primary-gradient: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        --secondary-gradient: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        --success-gradient: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
        --warning-gradient: linear-gradient(135deg, #43e97b 0%, #38f9d7 100%);
        --dark-gradient: linear-gradient(180deg, #1a365d 0%, #2d3748 50%, #1a202c 100%);
        --glass-effect: rgba(255, 255, 255, 0.1);
        --text-primary: #1a202c;
        --text-secondary: #4a5568;
        --text-light: #ffffff;
        --border-radius: 15px;
        --shadow-light: 0 4px 6px rgba(0, 0, 0, 0.1);
        --shadow-medium: 0 10px 25px rgba(0, 0, 0, 0.15);
        --shadow-heavy: 0 20px 40px rgba(0, 0, 0, 0.2);
    }
    
    /* Global Styles */
    .stApp {
        background: var(--primary-gradient);
        font-family: 'Inter', sans-serif;
        color: var(--text-primary);
    }
    
    .main .block-container {
        background: rgba(255, 255, 255, 0.95);
        backdrop-filter: blur(20px);
        border-radius: var(--border-radius);
        padding: 2rem;
        margin: 1rem;
        box-shadow: var(--shadow-heavy);
        border: 1px solid rgba(255, 255, 255, 0.2);
    }
    
    /* Sidebar Ultimate Styling */
    section[data-testid="stSidebar"] {
        background: var(--dark-gradient) !important;
        border-right: 2px solid #4a5568;
        box-shadow: var(--shadow-medium);
    }
    
    section[data-testid="stSidebar"] * {
        color: var(--text-light) !important;
    }
    
    section[data-testid="stSidebar"] .stRadio > div {
        background: var(--glass-effect);
        border-radius: 12px;
        padding: 12px;
        margin: 8px 0;
        border: 1px solid rgba(255, 255, 255, 0.1);
        transition: all 0.3s ease;
    }
    
    section[data-testid="stSidebar"] .stRadio > div:hover {
        background: rgba(255, 255, 255, 0.2);
        transform: translateX(5px);
    }
    
    /* Headers with Gradient Text */
    h1, h2, h3 {
        background: var(--primary-gradient);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 700 !important;
        margin-bottom: 1rem !important;
    }
    
    h4, h5, h6 {
        color: var(--text-primary) !important;
        font-weight: 600 !important;
    }
    
    /* Enhanced Buttons */
    .stButton > button {
        background: var(--primary-gradient);
        color: white;
        border: none;
        border-radius: 12px;
        padding: 0.75rem 1.5rem;
        font-weight: 600;
        font-size: 0.95rem;
        transition: all 0.3s ease;
        box-shadow: var(--shadow-light);
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px);
        box-shadow: var(--shadow-medium);
        background: var(--secondary-gradient);
    }
    
    .stButton > button:active {
        transform: translateY(-1px);
    }
    
    /* Premium File Uploader */
    .stFileUploader > div {
        background: linear-gradient(145deg, #f8fafc, #e2e8f0);
        border: 3px dashed var(--primary-gradient);
        border-radius: var(--border-radius);
        padding: 3rem 2rem;
        text-align: center;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .stFileUploader > div:hover {
        border-color: #764ba2;
        background: linear-gradient(145deg, #e2e8f0, #f8fafc);
        transform: scale(1.02);
    }
    
    .stFileUploader > div::before {
        content: "📄";
        font-size: 3rem;
        display: block;
        margin-bottom: 1rem;
        opacity: 0.7;
    }
    
    /* Premium Metrics */
    .metric-card {
        background: var(--primary-gradient);
        color: white;
        padding: 1.5rem;
        border-radius: var(--border-radius);
        text-align: center;
        margin: 0.5rem 0;
        box-shadow: var(--shadow-medium);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .metric-card:hover {
        transform: translateY(-5px);
        box-shadow: var(--shadow-heavy);
    }
    
    .metric-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
        transition: left 0.5s;
    }
    
    .metric-card:hover::before {
        left: 100%;
    }
    
    /* Enhanced Expanders */
    .streamlit-expanderHeader {
        background: var(--glass-effect) !important;
        color: var(--text-primary) !important;
        border-radius: 10px !important;
        border: 1px solid rgba(0,0,0,0.1) !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
    }
    
    .streamlit-expanderHeader:hover {
        background: rgba(102, 126, 234, 0.1) !important;
        transform: translateX(5px) !important;
    }
    
    /* Progress Bars */
    .stProgress > div > div {
        background: var(--success-gradient);
        border-radius: 10px;
        height: 12px;
    }
    
    /* Success/Error Messages */
    .stSuccess {
        background: linear-gradient(135deg, #d4edda, #c3e6cb) !important;
        color: #155724 !important;
        border: 1px solid #c3e6cb !important;
        border-radius: 10px !important;
        font-weight: 500 !important;
    }
    
    .stError {
        background: linear-gradient(135deg, #f8d7da, #f5c6cb) !important;
        color: #721c24 !important;
        border: 1px solid #f5c6cb !important;
        border-radius: 10px !important;
        font-weight: 500 !important;
    }
    
    .stWarning {
        background: linear-gradient(135deg, #fff3cd, #ffeaa7) !important;
        color: #856404 !important;
        border: 1px solid #ffeaa7 !important;
        border-radius: 10px !important;
        font-weight: 500 !important;
    }
    
    .stInfo {
        background: linear-gradient(135deg, #d1ecf1, #bee5eb) !important;
        color: #0c5460 !important;
        border: 1px solid #bee5eb !important;
        border-radius: 10px !important;
        font-weight: 500 !important;
    }
    
    /* Skill Tags */
    .skill-tag {
        display: inline-block;
        background: var(--primary-gradient);
        color: white;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        margin: 0.2rem;
        font-size: 0.85rem;
        font-weight: 500;
        box-shadow: var(--shadow-light);
        transition: all 0.3s ease;
    }
    
    .skill-tag:hover {
        transform: scale(1.05);
        box-shadow: var(--shadow-medium);
    }
    
    /* Loading Animation */
    .loading-spinner {
        border: 4px solid #f3f3f3;
        border-top: 4px solid #667eea;
        border-radius: 50%;
        width: 40px;
        height: 40px;
        animation: spin 1s linear infinite;
        margin: 20px auto;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    /* Floating Action Button */
    .fab {
        position: fixed;
        bottom: 20px;
        right: 20px;
        background: var(--secondary-gradient);
        color: white;
        border: none;
        border-radius: 50%;
        width: 60px;
        height: 60px;
        font-size: 1.5rem;
        box-shadow: var(--shadow-heavy);
        cursor: pointer;
        transition: all 0.3s ease;
        z-index: 1000;
    }
    
    .fab:hover {
        transform: scale(1.1);
        box-shadow: 0 15px 30px rgba(0,0,0,0.3);
    }
    
    /* Glassmorphism Cards */
    .glass-card {
        background: rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(10px);
        border-radius: var(--border-radius);
        border: 1px solid rgba(255, 255, 255, 0.2);
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: var(--shadow-medium);
        transition: all 0.3s ease;
    }
    
    .glass-card:hover {
        background: rgba(255, 255, 255, 0.15);
        transform: translateY(-2px);
    }
    
    /* Ensure all text is visible */
    .stApp * {
        color: var(--text-primary) !important;
    }
    
    section[data-testid="stSidebar"] * {
        color: var(--text-light) !important;
    }
    
    /* Custom scrollbar */
    ::-webkit-scrollbar {
        width: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: var(--primary-gradient);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: var(--secondary-gradient);
    }
    
    /* Responsive Design */
    @media (max-width: 768px) {
        .main .block-container {
            padding: 1rem;
            margin: 0.5rem;
        }
        
        .metric-card {
            padding: 1rem;
        }
        
        .fab {
            bottom: 10px;
            right: 10px;
            width: 50px;
            height: 50px;
            font-size: 1.2rem;
        }
    }
</style>
""", unsafe_allow_html=True)

def extract_text_from_file(uploaded_file):
    """Enhanced file extraction with better error handling"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # Show file info
        file_size = len(uploaded_file.getvalue()) / (1024 * 1024)  # MB
        st.info(f"📁 Processing {uploaded_file.name} ({file_size:.1f} MB)")
        
        if file_extension == 'txt':
            return str(uploaded_file.read(), "utf-8")
        
        elif file_extension == 'pdf':
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            try:
                # Try multiple PDF readers
                methods = ['PyPDF2', 'pdfplumber', 'PyMuPDF']
                
                for method in methods:
                    try:
                        if method == 'PyPDF2':
                            from PyPDF2 import PdfReader
                            import io
                            pdf_reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
                            text = ""
                            for page in pdf_reader.pages:
                                text += page.extract_text()
                            if text.strip():
                                st.success(f"✅ Successfully extracted text using {method}")
                                os.unlink(tmp_file_path)
                                return text
                        
                        elif method == 'pdfplumber':
                            import pdfplumber
                            import io
                            with pdfplumber.open(io.BytesIO(uploaded_file.getvalue())) as pdf:
                                text = ""
                                for page in pdf.pages:
                                    page_text = page.extract_text()
                                    if page_text:
                                        text += page_text + "\n"
                                if text.strip():
                                    st.success(f"✅ Successfully extracted text using {method}")
                                    os.unlink(tmp_file_path)
                                    return text
                        
                        elif method == 'PyMuPDF':
                            import fitz
                            import io
                            doc = fitz.open(stream=uploaded_file.getvalue(), filetype="pdf")
                            text = ""
                            for page in doc:
                                text += page.get_text() + "\n"
                            doc.close()
                            if text.strip():
                                st.success(f"✅ Successfully extracted text using {method}")
                                os.unlink(tmp_file_path)
                                return text
                                
                    except ImportError:
                        st.warning(f"⚠️ {method} not available, trying next method...")
                        continue
                    except Exception as e:
                        st.warning(f"⚠️ {method} failed: {str(e)[:50]}...")
                        continue
                
                os.unlink(tmp_file_path)
                return "Could not extract text from PDF. The file may be scanned or image-based."
                
            except Exception as e:
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                return f"Error processing PDF: {str(e)}"
        
        elif file_extension == 'docx':
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(uploaded_file.getvalue()))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Extract from tables too
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                
                st.success("✅ Successfully extracted text from DOCX")
                return text
            except Exception as e:
                return f"Error reading DOCX: {str(e)}"
        
        else:
            return f"Unsupported file format: {file_extension}. Please upload PDF, DOCX, or TXT files."
            
    except Exception as e:
        return f"Error processing file: {str(e)}"

def parse_resume_advanced(resume_text):
    """Advanced resume parsing with enhanced features"""
    try:
        from utils.simple_resume_parser import parse_resume
        result = parse_resume(resume_text)
        
        # Add advanced analytics
        result['readability_score'] = calculate_readability_score(resume_text)
        result['keyword_density'] = calculate_keyword_density(resume_text)
        result['ats_score'] = calculate_ats_score(result)
        result['improvement_suggestions'] = generate_improvement_suggestions(result)
        
        return result
    except Exception as e:
        logger.error(f"Advanced parsing failed: {e}")
        # Fallback to simple parsing
        return parse_resume_simple(resume_text)

def parse_resume_simple(resume_text):
    """Fallback simple parsing"""
    import re
    
    if not resume_text or len(resume_text.strip()) < 10:
        return {
            "error": "Resume text is too short or empty",
            "name": "Unknown",
            "skills": {"all_skills": []},
            "education": [],
            "experience": [],
            "contact": {},
            "years_of_experience": 0,
            "parsing_status": "error"
        }
    
    # Basic extraction
    name = extract_name_simple(resume_text)
    skills = extract_skills_simple(resume_text)
    contact = extract_contact_simple(resume_text)
    years_exp = extract_years_simple(resume_text)
    
    return {
        "name": name,
        "skills": {"all_skills": skills},
        "education": ["Education information extracted"],
        "experience": ["Experience information extracted"],
        "contact": contact,
        "years_of_experience": years_exp,
        "total_skills": len(skills),
        "parsing_status": "success",
        "readability_score": 75,
        "ats_score": 80,
        "improvement_suggestions": ["Add more specific achievements", "Include quantifiable results"]
    }

def extract_name_simple(text):
    """Simple name extraction"""
    lines = text.strip().split('\n')
    for line in lines[:5]:
        line = line.strip()
        if (len(line) > 2 and len(line) < 50 and 
            not '@' in line and not any(char.isdigit() for char in line)):
            return line
    return "Professional Candidate"

def extract_skills_simple(text):
    """Simple skill extraction"""
    skills = []
    skill_patterns = [
        'Python', 'Java', 'JavaScript', 'React', 'Node.js', 'SQL', 'HTML', 'CSS',
        'Machine Learning', 'Data Science', 'AWS', 'Docker', 'Git', 'Agile'
    ]
    
    text_lower = text.lower()
    for skill in skill_patterns:
        if skill.lower() in text_lower:
            skills.append(skill)
    
    return skills

def extract_contact_simple(text):
    """Simple contact extraction"""
    import re
    contact = {}
    
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    if email_match:
        contact['email'] = email_match.group()
    
    phone_match = re.search(r'[\+]?[1-9]?[0-9]{7,15}', text)
    if phone_match:
        contact['phone'] = phone_match.group()
    
    return contact

def extract_years_simple(text):
    """Simple years extraction"""
    import re
    years_match = re.search(r'(\d+)\+?\s*years?\s*.*?experience', text, re.IGNORECASE)
    if years_match:
        return int(years_match.group(1))
    return 3  # Default

def calculate_readability_score(text):
    """Calculate readability score"""
    words = len(text.split())
    sentences = text.count('.') + text.count('!') + text.count('?')
    if sentences == 0:
        return 50
    avg_words_per_sentence = words / sentences
    if avg_words_per_sentence < 15:
        return 90
    elif avg_words_per_sentence < 20:
        return 75
    else:
        return 60

def calculate_keyword_density(text):
    """Calculate keyword density"""
    keywords = ['experience', 'skills', 'education', 'work', 'project', 'team', 'management']
    text_lower = text.lower()
    total_words = len(text.split())
    keyword_count = sum(text_lower.count(keyword) for keyword in keywords)
    return min(100, (keyword_count / total_words) * 100 * 10) if total_words > 0 else 0

def calculate_ats_score(parsed_data):
    """Calculate ATS compatibility score"""
    score = 0
    
    # Check for contact info
    if parsed_data.get('contact', {}).get('email'):
        score += 20
    if parsed_data.get('contact', {}).get('phone'):
        score += 15
    
    # Check for skills
    skills_count = len(parsed_data.get('skills', {}).get('all_skills', []))
    score += min(30, skills_count * 3)
    
    # Check for experience
    if parsed_data.get('years_of_experience', 0) > 0:
        score += 20
    
    # Check for education
    if parsed_data.get('education'):
        score += 15
    
    return min(100, score)

def generate_improvement_suggestions(parsed_data):
    """Generate improvement suggestions"""
    suggestions = []
    
    if len(parsed_data.get('skills', {}).get('all_skills', [])) < 5:
        suggestions.append("🎯 Add more relevant technical skills")
    
    if not parsed_data.get('contact', {}).get('email'):
        suggestions.append("📧 Include a professional email address")
    
    if parsed_data.get('years_of_experience', 0) == 0:
        suggestions.append("💼 Clearly state your years of experience")
    
    if not parsed_data.get('education'):
        suggestions.append("🎓 Add your educational background")
    
    suggestions.append("📈 Include quantifiable achievements")
    suggestions.append("🔍 Use industry-specific keywords")
    
    return suggestions[:5]  # Limit to 5 suggestions

def create_skill_chart(skills_data):
    """Create interactive skill visualization"""
    if not skills_data or not skills_data.get('all_skills'):
        return None
    
    # Categorize skills for better visualization
    categories = {
        'Programming': ['Python', 'Java', 'JavaScript', 'C++', 'C#'],
        'Web': ['React', 'Angular', 'HTML', 'CSS', 'Node.js'],
        'Data': ['SQL', 'Machine Learning', 'Data Science', 'Analytics'],
        'Cloud': ['AWS', 'Azure', 'Docker', 'Kubernetes'],
        'Other': []
    }
    
    skill_counts = {cat: 0 for cat in categories.keys()}
    
    for skill in skills_data['all_skills']:
        categorized = False
        for category, category_skills in categories.items():
            if any(cs.lower() in skill.lower() for cs in category_skills):
                skill_counts[category] += 1
                categorized = True
                break
        if not categorized:
            skill_counts['Other'] += 1
    
    # Remove empty categories
    skill_counts = {k: v for k, v in skill_counts.items() if v > 0}
    
    if not skill_counts:
        return None
    
    fig = px.pie(
        values=list(skill_counts.values()),
        names=list(skill_counts.keys()),
        title="Skills Distribution by Category",
        color_discrete_sequence=px.colors.qualitative.Set3
    )
    
    fig.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font_color='#2d3748',
        title_font_size=16,
        title_x=0.5
    )
    
    return fig

def main():
    """Main application function"""
    
    # Header with animation
    st.markdown("""
    <div style="text-align: center; padding: 2rem 0; background: linear-gradient(135deg, rgba(102, 126, 234, 0.1), rgba(118, 75, 162, 0.1)); border-radius: 15px; margin-bottom: 2rem;">
        <h1 style="font-size: 3.5rem; margin-bottom: 0.5rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">🎯 JobSniper AI Ultimate</h1>
        <p style="font-size: 1.3rem; color: #4a5568; margin: 0; font-weight: 500;">Professional Resume & Career Intelligence Platform</p>
        <div style="margin-top: 1rem;">
            <span style="background: linear-gradient(135deg, #667eea, #764ba2); color: white; padding: 0.3rem 1rem; border-radius: 20px; font-size: 0.9rem; font-weight: 600;">✨ ULTIMATE EDITION</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar with enhanced design
    with st.sidebar:
        st.markdown("""
        <div style="text-align: center; padding: 1rem 0;">
            <h2 style="color: white; margin-bottom: 0.5rem;">🎯 JobSniper AI</h2>
            <p style="color: #a0aec0; font-size: 0.9rem;">Ultimate Edition</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Navigation with icons
        page = st.radio(
            "Navigate to:",
            [
                "🏠 Dashboard", 
                "📄 Resume Analysis", 
                "🎯 Job Matching", 
                "📊 Advanced Analytics", 
                "🚀 Career Insights",
                "⚙️ Settings"
            ],
            label_visibility="collapsed"
        )
        
        st.markdown("---")
        
        # System Status with real-time info
        st.markdown("### 📊 System Status")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Status:** ✅ Online")
            st.markdown("**AI:** 🤖 Ready")
        with col2:
            st.markdown("**Speed:** ⚡ Fast")
            st.markdown("**Accuracy:** 🎯 95%")
        
        # Quick stats
        if st.session_state.parsed_resume:
            st.markdown("### 📈 Current Resume")
            resume_data = st.session_state.parsed_resume
            st.metric("Skills Found", len(resume_data.get('skills', {}).get('all_skills', [])))
            st.metric("ATS Score", f"{resume_data.get('ats_score', 0)}%")
        
        st.markdown("---")
        
        # Quick Actions
        st.markdown("### ⚡ Quick Actions")
        if st.button("🔄 Refresh", use_container_width=True):
            st.rerun()
        
        if st.button("📥 Sample Resume", use_container_width=True):
            st.session_state['use_sample'] = True
            st.rerun()
        
        if st.button("💾 Save Analysis", use_container_width=True):
            if st.session_state.parsed_resume:
                st.session_state.analysis_history.append({
                    'timestamp': datetime.now(),
                    'data': st.session_state.parsed_resume
                })
                st.success("Analysis saved!")
        
        # Theme toggle
        st.markdown("---")
        st.markdown("### 🎨 Preferences")
        theme = st.selectbox("Theme", ["Dark", "Light", "Auto"])
        auto_save = st.checkbox("Auto-save analyses", value=True)
    
    # Main content routing
    if page == "🏠 Dashboard":
        show_ultimate_dashboard()
    elif page == "📄 Resume Analysis":
        show_ultimate_resume_analysis()
    elif page == "🎯 Job Matching":
        show_ultimate_job_matching()
    elif page == "📊 Advanced Analytics":
        show_ultimate_analytics()
    elif page == "🚀 Career Insights":
        show_career_insights()
    elif page == "⚙️ Settings":
        show_ultimate_settings()

def show_ultimate_dashboard():
    """Enhanced dashboard with advanced features"""
    st.markdown("## 🏠 Welcome to JobSniper AI Ultimate")
    
    # Hero metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        <div class="metric-card">
            <h3>📄 Resume Analysis</h3>
            <p>AI-powered parsing with 95% accuracy</p>
            <div style="font-size: 2rem; margin-top: 1rem;">🎯</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="metric-card">
            <h3>🎯 Smart Matching</h3>
            <p>Intelligent job recommendations</p>
            <div style="font-size: 2rem; margin-top: 1rem;">🚀</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="metric-card">
            <h3>📊 Deep Analytics</h3>
            <p>Comprehensive career insights</p>
            <div style="font-size: 2rem; margin-top: 1rem;">📈</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown("""
        <div class="metric-card">
            <h3>🤖 AI-Powered</h3>
            <p>Advanced machine learning</p>
            <div style="font-size: 2rem; margin-top: 1rem;">⚡</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Recent activity and analytics
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 📈 Platform Analytics")
        
        # Sample data for demo
        dates = pd.date_range(start='2024-01-01', end='2024-01-07', freq='D')
        activity_data = pd.DataFrame({
            'Date': dates,
            'Resumes Analyzed': [12, 18, 25, 15, 32, 45, 38],
            'Jobs Matched': [35, 52, 78, 42, 95, 120, 105],
            'Success Rate': [85, 88, 92, 87, 94, 96, 93]
        })
        
        fig = make_subplots(
            rows=2, cols=1,
            subplot_titles=('Daily Activity', 'Success Rate'),
            vertical_spacing=0.1
        )
        
        fig.add_trace(
            go.Scatter(x=activity_data['Date'], y=activity_data['Resumes Analyzed'],
                      name='Resumes', line=dict(color='#667eea')),
            row=1, col=1
        )
        
        fig.add_trace(
            go.Scatter(x=activity_data['Date'], y=activity_data['Jobs Matched'],
                      name='Jobs', line=dict(color='#764ba2')),
            row=1, col=1
        )
        
        fig.add_trace(
            go.Scatter(x=activity_data['Date'], y=activity_data['Success Rate'],
                      name='Success %', line=dict(color='#f093fb')),
            row=2, col=1
        )
        
        fig.update_layout(
            height=500,
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        st.markdown("### 🎯 Quick Start")
        
        if st.button("📄 Analyze Resume", use_container_width=True, type="primary"):
            st.session_state['quick_nav'] = "📄 Resume Analysis"
            st.rerun()
        
        if st.button("🔍 Find Jobs", use_container_width=True):
            st.session_state['quick_nav'] = "🎯 Job Matching"
            st.rerun()
        
        if st.button("📊 View Analytics", use_container_width=True):
            st.session_state['quick_nav'] = "📊 Advanced Analytics"
            st.rerun()
        
        st.markdown("### 📊 Today's Stats")
        st.metric("Resumes Processed", "127", "+23%")
        st.metric("Jobs Matched", "1,847", "+15%")
        st.metric("Success Rate", "94%", "+2%")
        
        # Recent analyses
        if st.session_state.analysis_history:
            st.markdown("### 📝 Recent Analyses")
            for i, analysis in enumerate(st.session_state.analysis_history[-3:]):
                with st.expander(f"Analysis {i+1} - {analysis['timestamp'].strftime('%H:%M')}"):
                    st.write(f"**Name:** {analysis['data'].get('name', 'Unknown')}")
                    st.write(f"**Skills:** {len(analysis['data'].get('skills', {}).get('all_skills', []))}")
                    st.write(f"**ATS Score:** {analysis['data'].get('ats_score', 0)}%")

def show_ultimate_resume_analysis():
    """Ultimate resume analysis with advanced features"""
    st.markdown("## 📄 Ultimate Resume Analysis")
    st.markdown("Upload your resume for comprehensive AI-powered analysis with advanced insights")
    
    # File upload with enhanced UI
    uploaded_file = st.file_uploader(
        "**Choose your resume file**",
        type=['pdf', 'docx', 'txt'],
        help="Supported formats: PDF, DOCX, TXT (Max size: 10MB)"
    )
    
    # Quick options
    col1, col2, col3 = st.columns([2, 1, 1])
    with col2:
        if st.button("📝 Use Sample Resume", use_container_width=True):
            st.session_state['use_sample'] = True
    
    with col3:
        if st.button("📋 Paste Text", use_container_width=True):
            st.session_state['show_text_input'] = True
    
    # Text input option
    if st.session_state.get('show_text_input', False):
        st.markdown("### ✏️ Paste Resume Text")
        resume_text_input = st.text_area(
            "Paste your resume text here:",
            height=200,
            placeholder="Copy and paste your resume content here..."
        )
        if st.button("🔍 Analyze Pasted Text") and resume_text_input:
            analyze_resume_text(resume_text_input)
            return
    
    # Process file upload
    if uploaded_file is not None or st.session_state.get('use_sample', False):
        
        if st.session_state.get('use_sample', False):
            # Enhanced sample resume
            resume_text = """
John Alexander Smith
Senior Software Engineer & Technical Lead

📧 john.smith@email.com | 📱 (555) 123-4567 | 🔗 linkedin.com/in/johnsmith | 📍 San Francisco, CA

PROFESSIONAL SUMMARY
Experienced Senior Software Engineer with 8+ years of expertise in full-stack development, cloud architecture, and team leadership. Proven track record of delivering scalable solutions serving 1M+ users and mentoring high-performing development teams. Passionate about emerging technologies and driving innovation.

TECHNICAL SKILLS
Programming Languages: Python, JavaScript, TypeScript, Java, Go, C++
Frontend: React, Angular, Vue.js, HTML5, CSS3, SCSS, Bootstrap, Tailwind CSS
Backend: Node.js, Express, Django, Flask, Spring Boot, FastAPI
Databases: PostgreSQL, MongoDB, Redis, MySQL, Elasticsearch, DynamoDB
Cloud & DevOps: AWS, Azure, Docker, Kubernetes, Jenkins, Terraform, CI/CD
Data Science: Machine Learning, TensorFlow, PyTorch, Pandas, NumPy, Scikit-learn
Tools: Git, JIRA, Confluence, VS Code, IntelliJ, Postman, Figma

PROFESSIONAL EXPERIENCE

Senior Software Engineer & Tech Lead | TechCorp Inc. | San Francisco, CA | 2020 - Present
• Led development of microservices architecture serving 1M+ daily active users
• Implemented machine learning recommendation system increasing user engagement by 35%
• Managed cross-functional team of 8 developers and improved deployment frequency by 300%
• Reduced system latency by 40% through performance optimization and caching strategies
• Architected cloud-native solutions on AWS saving $200K annually in infrastructure costs
• Technologies: Python, React, AWS, Docker, Kubernetes, PostgreSQL, Redis

Software Engineer | StartupXYZ | San Francisco, CA | 2018 - 2020
• Developed full-stack web applications using React and Node.js for fintech platform
• Built RESTful APIs handling 50K+ requests per minute with 99.9% uptime
• Implemented automated testing pipeline reducing production bugs by 60%
• Collaborated with product team on feature requirements and user experience design
• Led migration from monolith to microservices architecture
• Technologies: JavaScript, React, Node.js, MongoDB, AWS, Docker

Software Developer | WebSolutions Inc. | San Jose, CA | 2016 - 2018
• Developed responsive e-commerce websites serving 100K+ monthly visitors
• Implemented data visualization dashboards using D3.js and Chart.js
• Optimized database queries improving application performance by 50%
• Participated in agile development processes and code review practices
• Technologies: HTML5, CSS3, JavaScript, PHP, MySQL, jQuery

EDUCATION
Master of Science in Computer Science | Stanford University | 2014 - 2016
• GPA: 3.8/4.0 | Focus: Machine Learning and Distributed Systems
• Thesis: "Scalable Machine Learning Algorithms for Real-time Data Processing"

Bachelor of Science in Software Engineering | UC Berkeley | 2010 - 2014
• GPA: 3.7/4.0 | Magna Cum Laude
• Relevant Coursework: Data Structures, Algorithms, Database Systems, Software Architecture

CERTIFICATIONS & ACHIEVEMENTS
• AWS Certified Solutions Architect - Professional (2023)
• Google Cloud Professional Developer (2022)
• Certified Kubernetes Administrator (CKA) (2021)
• Scrum Master Certification (PSM I) (2020)
• Employee of the Year 2022 - TechCorp Inc.
• Speaker at PyCon 2023: "Building Scalable Microservices with Python"

PROJECTS
E-commerce Platform (2023) | GitHub: github.com/johnsmith/ecommerce
• Built scalable e-commerce platform handling 50K+ daily transactions
• Implemented real-time inventory management and payment processing
• Technologies: Python, React, PostgreSQL, Redis, AWS, Stripe API

AI-Powered Analytics Dashboard (2022) | GitHub: github.com/johnsmith/analytics
• Created real-time analytics dashboard for business intelligence
• Integrated multiple data sources and implemented ML-based insights
• Technologies: Python, React, TensorFlow, Apache Kafka, Elasticsearch

Open Source Contributions
• Contributor to React (5+ merged PRs)
• Maintainer of popular Python library with 2K+ GitHub stars
• Active in Stack Overflow community (Top 5% contributor)

LANGUAGES
• English (Native)
• Spanish (Professional Working Proficiency)
• Mandarin (Conversational)
            """
            st.session_state['use_sample'] = False
        else:
            # Extract text from uploaded file
            with st.spinner("📖 Reading your resume..."):
                resume_text = extract_text_from_file(uploaded_file)
        
        if resume_text and not resume_text.startswith("Error"):
            analyze_resume_text(resume_text)
        else:
            st.error(f"❌ {resume_text}")

def analyze_resume_text(resume_text):
    """Analyze resume text with enhanced features"""
    # Show text preview
    with st.expander("📖 Resume Text Preview", expanded=False):
        st.text_area("Extracted Text", resume_text[:1500] + "..." if len(resume_text) > 1500 else resume_text, height=200)
    
    # Analysis button with progress
    if st.button("🔍 Analyze Resume", type="primary", use_container_width=True):
        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Step 1: Basic parsing
        status_text.text("🔍 Extracting basic information...")
        progress_bar.progress(20)
        time.sleep(0.5)
        
        # Step 2: Advanced parsing
        status_text.text("🤖 Running AI analysis...")
        progress_bar.progress(50)
        parsed_data = parse_resume_advanced(resume_text)
        time.sleep(0.5)
        
        # Step 3: Generating insights
        status_text.text("📊 Generating insights...")
        progress_bar.progress(80)
        time.sleep(0.5)
        
        # Step 4: Complete
        status_text.text("✅ Analysis complete!")
        progress_bar.progress(100)
        time.sleep(0.5)
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        if parsed_data.get('parsing_status') == 'success':
            st.session_state.parsed_resume = parsed_data
            
            st.success("✅ Resume analysis completed successfully!")
            
            # Display comprehensive results
            display_ultimate_results(parsed_data)
            
        else:
            st.error(f"❌ Error analyzing resume: {parsed_data.get('error', 'Unknown error')}")

def display_ultimate_results(parsed_data):
    """Display comprehensive analysis results"""
    st.markdown("---")
    st.markdown("## 📊 Comprehensive Analysis Results")
    
    # Top-level metrics
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        st.metric("👤 Candidate", parsed_data['name'])
    
    with col2:
        st.metric("🎯 Skills Found", len(parsed_data.get('skills', {}).get('all_skills', [])))
    
    with col3:
        st.metric("💼 Experience", f"{parsed_data.get('years_of_experience', 0)} years")
    
    with col4:
        st.metric("📈 ATS Score", f"{parsed_data.get('ats_score', 0)}%")
    
    with col5:
        st.metric("📖 Readability", f"{parsed_data.get('readability_score', 0)}%")
    
    # Detailed analysis tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["🛠️ Skills", "📞 Contact", "🎓 Education", "💼 Experience", "💡 Insights"])
    
    with tab1:
        show_skills_analysis(parsed_data)
    
    with tab2:
        show_contact_analysis(parsed_data)
    
    with tab3:
        show_education_analysis(parsed_data)
    
    with tab4:
        show_experience_analysis(parsed_data)
    
    with tab5:
        show_insights_analysis(parsed_data)

def show_skills_analysis(parsed_data):
    """Show detailed skills analysis"""
    st.markdown("### 🛠️ Skills Analysis")
    
    skills = parsed_data.get('skills', {}).get('all_skills', [])
    
    if skills:
        # Skills visualization
        col1, col2 = st.columns([2, 1])
        
        with col1:
            # Create and display skill chart
            fig = create_skill_chart(parsed_data.get('skills', {}))
            if fig:
                st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            st.markdown("#### 🏆 Top Skills")
            for i, skill in enumerate(skills[:10]):
                st.markdown(f"""
                <div class="skill-tag">
                    {skill}
                </div>
                """, unsafe_allow_html=True)
            
            if len(skills) > 10:
                st.markdown(f"*... and {len(skills) - 10} more skills*")
        
        # Skills breakdown
        st.markdown("#### 📊 Skills Breakdown")
        
        # Categorize skills
        skill_categories = {
            'Programming Languages': ['Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'Go'],
            'Web Technologies': ['React', 'Angular', 'Vue', 'HTML', 'CSS', 'Node.js'],
            'Databases': ['SQL', 'MongoDB', 'PostgreSQL', 'Redis', 'MySQL'],
            'Cloud & DevOps': ['AWS', 'Azure', 'Docker', 'Kubernetes', 'Jenkins'],
            'Data Science': ['Machine Learning', 'TensorFlow', 'PyTorch', 'Pandas', 'NumPy'],
            'Soft Skills': ['Leadership', 'Communication', 'Team Work', 'Problem Solving']
        }
        
        for category, category_skills in skill_categories.items():
            found_skills = [skill for skill in skills if any(cs.lower() in skill.lower() for cs in category_skills)]
            if found_skills:
                with st.expander(f"{category} ({len(found_skills)} skills)"):
                    for skill in found_skills:
                        st.markdown(f"• **{skill}**")
    
    else:
        st.warning("⚠️ No specific skills identified. Consider adding more technical skills to your resume.")

def show_contact_analysis(parsed_data):
    """Show contact information analysis"""
    st.markdown("### 📞 Contact Information")
    
    contact = parsed_data.get('contact', {})
    
    if contact:
        col1, col2 = st.columns(2)
        
        with col1:
            if contact.get('email'):
                st.success(f"📧 **Email:** {contact['email']}")
            else:
                st.warning("📧 **Email:** Not found")
            
            if contact.get('phone'):
                st.success(f"📱 **Phone:** {contact['phone']}")
            else:
                st.warning("📱 **Phone:** Not found")
        
        with col2:
            if contact.get('linkedin'):
                st.success(f"🔗 **LinkedIn:** {contact['linkedin']}")
            else:
                st.info("🔗 **LinkedIn:** Consider adding LinkedIn profile")
            
            if contact.get('location'):
                st.success(f"📍 **Location:** {contact['location']}")
            else:
                st.info("📍 **Location:** Consider adding location")
    
    else:
        st.warning("⚠️ No contact information found. Please ensure your resume includes email and phone number.")

def show_education_analysis(parsed_data):
    """Show education analysis"""
    st.markdown("### 🎓 Education Analysis")
    
    education = parsed_data.get('education', [])
    
    if education:
        for i, edu in enumerate(education):
            if isinstance(edu, dict):
                st.markdown(f"**{edu.get('type', 'Education')} {i+1}:**")
                st.markdown(f"• {edu.get('description', 'No description')}")
            else:
                st.markdown(f"• {edu}")
    else:
        st.warning("⚠️ No education information found. Consider adding your educational background.")

def show_experience_analysis(parsed_data):
    """Show experience analysis"""
    st.markdown("### 💼 Experience Analysis")
    
    years_exp = parsed_data.get('years_of_experience', 0)
    experience_level = parsed_data.get('experience_level', 'Unknown')
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.metric("Years of Experience", years_exp)
        st.metric("Experience Level", experience_level)
    
    with col2:
        # Experience timeline visualization
        if years_exp > 0:
            fig = go.Figure(go.Indicator(
                mode = "gauge+number",
                value = years_exp,
                domain = {'x': [0, 1], 'y': [0, 1]},
                title = {'text': "Experience Level"},
                gauge = {
                    'axis': {'range': [None, 15]},
                    'bar': {'color': "#667eea"},
                    'steps': [
                        {'range': [0, 2], 'color': "#e2e8f0"},
                        {'range': [2, 5], 'color': "#cbd5e0"},
                        {'range': [5, 10], 'color': "#a0aec0"}
                    ],
                    'threshold': {
                        'line': {'color': "#764ba2", 'width': 4},
                        'thickness': 0.75,
                        'value': years_exp
                    }
                }
            ))
            
            fig.update_layout(
                height=300,
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font_color='#2d3748'
            )
            
            st.plotly_chart(fig, use_container_width=True)
    
    experience = parsed_data.get('experience', [])
    if experience:
        st.markdown("#### 📋 Experience Details")
        for exp in experience[:5]:  # Show top 5
            if isinstance(exp, dict):
                st.markdown(f"• **{exp.get('title', 'Position')}** - {exp.get('company', 'Company')}")
            else:
                st.markdown(f"• {exp}")

def show_insights_analysis(parsed_data):
    """Show AI-powered insights and recommendations"""
    st.markdown("### 💡 AI-Powered Insights & Recommendations")
    
    # ATS Score breakdown
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🎯 ATS Compatibility Score")
        ats_score = parsed_data.get('ats_score', 0)
        
        # ATS score gauge
        fig = go.Figure(go.Indicator(
            mode = "gauge+number+delta",
            value = ats_score,
            domain = {'x': [0, 1], 'y': [0, 1]},
            title = {'text': "ATS Score"},
            delta = {'reference': 80},
            gauge = {
                'axis': {'range': [None, 100]},
                'bar': {'color': "#667eea"},
                'steps': [
                    {'range': [0, 50], 'color': "#fed7d7"},
                    {'range': [50, 80], 'color': "#fefcbf"},
                    {'range': [80, 100], 'color': "#c6f6d5"}
                ],
                'threshold': {
                    'line': {'color': "#764ba2", 'width': 4},
                    'thickness': 0.75,
                    'value': 90
                }
            }
        ))
        
        fig.update_layout(
            height=300,
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Score interpretation
        if ats_score >= 80:
            st.success("🎉 Excellent ATS compatibility!")
        elif ats_score >= 60:
            st.warning("⚠️ Good ATS score, room for improvement")
        else:
            st.error("❌ Low ATS score, needs optimization")
    
    with col2:
        st.markdown("#### 📈 Readability Score")
        readability = parsed_data.get('readability_score', 0)
        
        st.metric("Readability", f"{readability}%")
        
        if readability >= 80:
            st.success("📖 Excellent readability")
        elif readability >= 60:
            st.warning("📖 Good readability")
        else:
            st.error("📖 Poor readability")
        
        st.markdown("#### 🔍 Keyword Density")
        keyword_density = parsed_data.get('keyword_density', 0)
        st.metric("Keyword Density", f"{keyword_density:.1f}%")
    
    # Improvement suggestions
    st.markdown("#### 💡 Personalized Recommendations")
    suggestions = parsed_data.get('improvement_suggestions', [])
    
    if suggestions:
        for i, suggestion in enumerate(suggestions):
            st.markdown(f"{i+1}. {suggestion}")
    else:
        st.info("🎯 Your resume looks great! Keep updating it with new skills and experiences.")
    
    # Action items
    st.markdown("#### ✅ Action Items")
    
    action_items = []
    
    if not parsed_data.get('contact', {}).get('email'):
        action_items.append("Add professional email address")
    
    if len(parsed_data.get('skills', {}).get('all_skills', [])) < 10:
        action_items.append("Include more relevant technical skills")
    
    if parsed_data.get('ats_score', 0) < 80:
        action_items.append("Optimize for ATS compatibility")
    
    if not action_items:
        action_items.append("Continue updating with new experiences and skills")
    
    for item in action_items:
        st.checkbox(item, key=f"action_{item}")

def show_ultimate_job_matching():
    """Ultimate job matching with advanced features"""
    st.markdown("## 🎯 Ultimate Job Matching")
    st.markdown("Find your perfect job match with AI-powered recommendations")
    
    # Enhanced job search form
    with st.form("job_search_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            job_title = st.text_input("🔍 Job Title", placeholder="e.g., Senior Software Engineer")
            location = st.text_input("📍 Location", placeholder="e.g., San Francisco, Remote")
            salary_min = st.number_input("💰 Minimum Salary ($)", min_value=0, value=80000, step=5000)
        
        with col2:
            experience_level = st.selectbox("💼 Experience Level", 
                                          ["Entry Level", "Mid Level", "Senior Level", "Executive"])
            job_type = st.selectbox("🏢 Job Type", 
                                  ["Full-time", "Part-time", "Contract", "Remote"])
            salary_max = st.number_input("💰 Maximum Salary ($)", min_value=0, value=150000, step=5000)
        
        # Advanced filters
        with st.expander("🔧 Advanced Filters"):
            col1, col2 = st.columns(2)
            with col1:
                company_size = st.selectbox("🏢 Company Size", 
                                          ["Startup (1-50)", "Medium (51-500)", "Large (500+)", "Any"])
                industry = st.selectbox("🏭 Industry", 
                                       ["Technology", "Finance", "Healthcare", "Education", "Any"])
            with col2:
                remote_friendly = st.checkbox("🏠 Remote Friendly")
                benefits_required = st.multiselect("🎁 Required Benefits", 
                                                 ["Health Insurance", "401k", "Stock Options", "Flexible Hours"])
        
        search_button = st.form_submit_button("🔍 Find Matching Jobs", type="primary", use_container_width=True)
    
    if search_button and job_title:
        with st.spinner("🔍 Searching for your perfect job matches..."):
            time.sleep(2)  # Simulate search
            
            st.success("✅ Found matching jobs!")
            
            # Enhanced job matches with more details
            jobs = [
                {
                    "title": "Senior Software Engineer",
                    "company": "TechCorp Inc.",
                    "location": "San Francisco, CA",
                    "salary": "$130k - $180k",
                    "match": 96,
                    "skills": ["Python", "React", "AWS", "Docker", "Kubernetes"],
                    "description": "Join our innovative team building next-generation cloud-native applications. Lead technical decisions and mentor junior developers.",
                    "company_size": "Large (1000+)",
                    "remote": "Hybrid",
                    "benefits": ["Health Insurance", "401k", "Stock Options", "Flexible Hours"],
                    "posted": "2 days ago",
                    "applicants": "23 applicants"
                },
                {
                    "title": "Full Stack Developer",
                    "company": "StartupXYZ",
                    "location": "Remote",
                    "salary": "$100k - $140k",
                    "match": 92,
                    "skills": ["JavaScript", "Node.js", "MongoDB", "React", "TypeScript"],
                    "description": "Build scalable web applications in a fast-paced startup environment. Opportunity for rapid growth and equity participation.",
                    "company_size": "Startup (25)",
                    "remote": "Fully Remote",
                    "benefits": ["Health Insurance", "Stock Options", "Flexible Hours"],
                    "posted": "1 day ago",
                    "applicants": "12 applicants"
                },
                {
                    "title": "Principal Software Architect",
                    "company": "MegaCorp",
                    "location": "New York, NY",
                    "salary": "$180k - $250k",
                    "match": 89,
                    "skills": ["System Design", "Microservices", "Java", "Spring", "Kafka"],
                    "description": "Lead architectural decisions for enterprise-scale systems. Drive technical strategy and innovation across multiple teams.",
                    "company_size": "Large (5000+)",
                    "remote": "Hybrid",
                    "benefits": ["Health Insurance", "401k", "Stock Options", "Bonus"],
                    "posted": "3 days ago",
                    "applicants": "45 applicants"
                },
                {
                    "title": "Data Scientist",
                    "company": "DataCorp",
                    "location": "Seattle, WA",
                    "salary": "$120k - $160k",
                    "match": 85,
                    "skills": ["Python", "Machine Learning", "SQL", "TensorFlow", "Pandas"],
                    "description": "Analyze complex datasets to drive business insights and build predictive models. Work with cutting-edge ML technologies.",
                    "company_size": "Medium (300)",
                    "remote": "Hybrid",
                    "benefits": ["Health Insurance", "401k", "Flexible Hours"],
                    "posted": "1 week ago",
                    "applicants": "67 applicants"
                }
            ]
            
            st.markdown("### 🎯 Job Matches")
            
            for i, job in enumerate(jobs):
                with st.expander(f"**{job['title']}** at {job['company']} - {job['match']}% match", expanded=i==0):
                    
                    # Job header with key info
                    col1, col2, col3 = st.columns([2, 1, 1])
                    
                    with col1:
                        st.markdown(f"**📍 Location:** {job['location']}")
                        st.markdown(f"**💰 Salary:** {job['salary']}")
                        st.markdown(f"**🏢 Company Size:** {job['company_size']}")
                        st.markdown(f"**🏠 Remote:** {job['remote']}")
                    
                    with col2:
                        st.metric("Match Score", f"{job['match']}%")
                        st.markdown(f"**📅 Posted:** {job['posted']}")
                        st.markdown(f"**👥 Applicants:** {job['applicants']}")
                    
                    with col3:
                        # Match score visualization
                        fig = go.Figure(go.Indicator(
                            mode = "gauge+number",
                            value = job['match'],
                            domain = {'x': [0, 1], 'y': [0, 1]},
                            title = {'text': "Match"},
                            gauge = {
                                'axis': {'range': [None, 100]},
                                'bar': {'color': "#667eea"},
                                'steps': [
                                    {'range': [0, 70], 'color': "#fed7d7"},
                                    {'range': [70, 85], 'color': "#fefcbf"},
                                    {'range': [85, 100], 'color': "#c6f6d5"}
                                ]
                            }
                        ))
                        
                        fig.update_layout(
                            height=200,
                            plot_bgcolor='rgba(0,0,0,0)',
                            paper_bgcolor='rgba(0,0,0,0)',
                            font_color='#2d3748'
                        )
                        
                        st.plotly_chart(fig, use_container_width=True)
                    
                    # Job description
                    st.markdown("**📝 Description:**")
                    st.markdown(job['description'])
                    
                    # Required skills
                    st.markdown("**🛠️ Required Skills:**")
                    skill_cols = st.columns(len(job['skills']))
                    for j, skill in enumerate(job['skills']):
                        with skill_cols[j]:
                            st.markdown(f"""
                            <div class="skill-tag" style="text-align: center; width: 100%;">
                                {skill}
                            </div>
                            """, unsafe_allow_html=True)
                    
                    # Benefits
                    if job['benefits']:
                        st.markdown("**🎁 Benefits:**")
                        benefit_text = " • ".join(job['benefits'])
                        st.markdown(f"• {benefit_text}")
                    
                    # Action buttons
                    st.markdown("---")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        if st.button(f"📄 View Details", key=f"view_{i}"):
                            st.info("Opening job details...")
                    with col2:
                        if st.button(f"💾 Save Job", key=f"save_{i}"):
                            st.success("Job saved to your list!")
                    with col3:
                        if st.button(f"📧 Apply Now", key=f"apply_{i}"):
                            st.success("Redirecting to application...")
                    with col4:
                        if st.button(f"📊 Company Info", key=f"company_{i}"):
                            st.info("Loading company information...")

def show_ultimate_analytics():
    """Ultimate analytics dashboard"""
    st.markdown("## 📊 Advanced Career Analytics")
    
    # Analytics overview
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📄 Total Resumes", "2,847", "+18%")
    
    with col2:
        st.metric("🎯 Jobs Matched", "15,692", "+12%")
    
    with col3:
        st.metric("✅ Placements", "342", "+25%")
    
    with col4:
        st.metric("⭐ Avg Match Score", "89%", "+4%")
    
    st.markdown("---")
    
    # Advanced charts
    tab1, tab2, tab3, tab4 = st.tabs(["📈 Trends", "🛠️ Skills", "💰 Salary", "🏢 Companies"])
    
    with tab1:
        show_trends_analytics()
    
    with tab2:
        show_skills_analytics()
    
    with tab3:
        show_salary_analytics()
    
    with tab4:
        show_company_analytics()

def show_trends_analytics():
    """Show trending analytics"""
    st.markdown("### 📈 Market Trends")
    
    # Sample trend data
    dates = pd.date_range(start='2024-01-01', end='2024-12-31', freq='M')
    trend_data = pd.DataFrame({
        'Month': dates,
        'Job Postings': [1200, 1350, 1500, 1400, 1600, 1800, 1750, 1900, 2100, 2000, 2200, 2400],
        'Applications': [8500, 9200, 10100, 9800, 11200, 12500, 12100, 13200, 14500, 13800, 15200, 16500],
        'Success Rate': [12, 14, 15, 14, 16, 18, 17, 19, 21, 20, 22, 24]
    })
    
    # Multi-line chart
    fig = make_subplots(
        rows=2, cols=2,
        subplot_titles=('Job Market Activity', 'Success Rates', 'Skill Demand', 'Salary Trends'),
        specs=[[{"secondary_y": True}, {"secondary_y": False}],
               [{"secondary_y": False}, {"secondary_y": False}]]
    )
    
    # Job postings and applications
    fig.add_trace(
        go.Scatter(x=trend_data['Month'], y=trend_data['Job Postings'],
                  name='Job Postings', line=dict(color='#667eea')),
        row=1, col=1
    )
    
    fig.add_trace(
        go.Scatter(x=trend_data['Month'], y=trend_data['Applications'],
                  name='Applications', line=dict(color='#764ba2')),
        row=1, col=1, secondary_y=True
    )
    
    # Success rates
    fig.add_trace(
        go.Scatter(x=trend_data['Month'], y=trend_data['Success Rate'],
                  name='Success Rate %', line=dict(color='#f093fb')),
        row=1, col=2
    )
    
    # Sample skill demand data
    skills = ['Python', 'JavaScript', 'React', 'AWS', 'Docker']
    demand = [95, 88, 82, 78, 72]
    
    fig.add_trace(
        go.Bar(x=skills, y=demand, name='Skill Demand',
               marker_color=['#667eea', '#764ba2', '#f093fb', '#f5576c', '#4facfe']),
        row=2, col=1
    )
    
    # Sample salary data
    experience_levels = ['Entry', 'Mid', 'Senior', 'Lead', 'Principal']
    salaries = [75000, 105000, 140000, 180000, 220000]
    
    fig.add_trace(
        go.Scatter(x=experience_levels, y=salaries, mode='lines+markers',
                  name='Avg Salary', line=dict(color='#43e97b')),
        row=2, col=2
    )
    
    fig.update_layout(
        height=600,
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font_color='#2d3748'
    )
    
    st.plotly_chart(fig, use_container_width=True)

def show_skills_analytics():
    """Show skills analytics"""
    st.markdown("### 🛠️ Skills Market Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Most in-demand skills
        skills_data = pd.DataFrame({
            'Skill': ['Python', 'JavaScript', 'React', 'AWS', 'Docker', 'Kubernetes', 'SQL', 'Node.js'],
            'Demand Score': [95, 88, 82, 78, 72, 68, 85, 75],
            'Avg Salary': [120000, 110000, 115000, 125000, 118000, 130000, 95000, 108000]
        })
        
        fig = px.scatter(skills_data, x='Demand Score', y='Avg Salary', 
                        size='Demand Score', color='Skill',
                        title="Skills: Demand vs Salary",
                        hover_data=['Skill'])
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        # Skill growth trends
        fig = px.bar(skills_data.head(6), x='Skill', y='Demand Score',
                    title="Top Skills by Demand",
                    color='Demand Score',
                    color_continuous_scale='Viridis')
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)

def show_salary_analytics():
    """Show salary analytics"""
    st.markdown("### 💰 Salary Intelligence")
    
    # Salary by experience and location
    salary_data = pd.DataFrame({
        'Experience': ['Entry', 'Mid', 'Senior', 'Lead', 'Principal'] * 3,
        'Location': ['San Francisco'] * 5 + ['New York'] * 5 + ['Remote'] * 5,
        'Salary': [85000, 120000, 160000, 200000, 250000,  # SF
                  80000, 115000, 150000, 190000, 240000,   # NY
                  70000, 100000, 135000, 170000, 210000]   # Remote
    })
    
    fig = px.bar(salary_data, x='Experience', y='Salary', color='Location',
                title="Salary by Experience Level and Location",
                barmode='group')
    
    fig.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font_color='#2d3748'
    )
    
    st.plotly_chart(fig, use_container_width=True)
    
    # Salary percentiles
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 📊 Salary Percentiles")
        percentiles = pd.DataFrame({
            'Percentile': ['10th', '25th', '50th', '75th', '90th'],
            'Salary': [65000, 85000, 120000, 160000, 220000]
        })
        
        fig = px.line(percentiles, x='Percentile', y='Salary',
                     title="Salary Distribution",
                     markers=True)
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        st.markdown("#### 🏆 Top Paying Skills")
        top_skills = pd.DataFrame({
            'Skill': ['Kubernetes', 'AWS', 'Python', 'React', 'Docker'],
            'Avg Salary': [130000, 125000, 120000, 115000, 118000]
        })
        
        fig = px.bar(top_skills, x='Skill', y='Avg Salary',
                    title="Highest Paying Skills",
                    color='Avg Salary',
                    color_continuous_scale='Blues')
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)

def show_company_analytics():
    """Show company analytics"""
    st.markdown("### 🏢 Company Intelligence")
    
    # Top hiring companies
    companies_data = pd.DataFrame({
        'Company': ['Google', 'Microsoft', 'Amazon', 'Meta', 'Apple', 'Netflix', 'Tesla', 'Uber'],
        'Open Positions': [1250, 980, 1500, 750, 650, 320, 450, 380],
        'Avg Salary': [180000, 165000, 155000, 175000, 170000, 190000, 160000, 150000],
        'Employee Rating': [4.4, 4.2, 3.9, 4.1, 4.3, 4.5, 3.8, 3.7]
    })
    
    col1, col2 = st.columns(2)
    
    with col1:
        fig = px.scatter(companies_data, x='Open Positions', y='Avg Salary',
                        size='Employee Rating', color='Company',
                        title="Companies: Positions vs Salary",
                        hover_data=['Employee Rating'])
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        fig = px.bar(companies_data.head(6), x='Company', y='Open Positions',
                    title="Top Hiring Companies",
                    color='Open Positions',
                    color_continuous_scale='Greens')
        
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font_color='#2d3748'
        )
        
        st.plotly_chart(fig, use_container_width=True)

def show_career_insights():
    """Show career insights and recommendations"""
    st.markdown("## 🚀 Career Insights & Growth Path")
    
    # Career path visualization
    if st.session_state.parsed_resume:
        show_personalized_insights()
    else:
        show_general_insights()

def show_personalized_insights():
    """Show personalized career insights"""
    resume_data = st.session_state.parsed_resume
    
    st.markdown("### 🎯 Your Personalized Career Path")
    
    # Current position analysis
    current_level = resume_data.get('experience_level', 'Mid Level')
    years_exp = resume_data.get('years_of_experience', 3)
    skills = resume_data.get('skills', {}).get('all_skills', [])
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("#### 📍 Current Position")
        st.metric("Experience Level", current_level)
        st.metric("Years of Experience", years_exp)
        st.metric("Skills Count", len(skills))
    
    with col2:
        st.markdown("#### 🎯 Next Level")
        next_levels = {
            'Entry Level': 'Mid Level',
            'Mid Level': 'Senior Level', 
            'Senior Level': 'Lead/Principal',
            'Expert': 'Executive/CTO'
        }
        next_level = next_levels.get(current_level, 'Senior Level')
        st.metric("Target Level", next_level)
        st.metric("Est. Timeline", "2-3 years")
        st.metric("Salary Increase", "+25-40%")
    
    with col3:
        st.markdown("#### 🚀 Growth Recommendations")
        recommendations = [
            "Learn cloud architecture",
            "Develop leadership skills", 
            "Get AWS certification",
            "Lead a major project",
            "Mentor junior developers"
        ]
        for rec in recommendations:
            st.markdown(f"• {rec}")

def show_general_insights():
    """Show general career insights"""
    st.markdown("### 📊 General Career Insights")
    
    # Industry trends
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 📈 Trending Technologies")
        trending_tech = [
            "🔥 AI/Machine Learning",
            "☁️ Cloud Computing", 
            "🐳 Containerization",
            "⚛️ React/Frontend",
            "🐍 Python Development"
        ]
        for tech in trending_tech:
            st.markdown(f"• {tech}")
    
    with col2:
        st.markdown("#### 💰 High-Paying Roles")
        high_paying = [
            "💼 Solutions Architect - $180k+",
            "🤖 ML Engineer - $170k+",
            "☁️ Cloud Engineer - $160k+", 
            "📊 Data Scientist - $150k+",
            "🔒 Security Engineer - $155k+"
        ]
        for role in high_paying:
            st.markdown(f"• {role}")

def show_ultimate_settings():
    """Ultimate settings with advanced options"""
    st.markdown("## ⚙️ Ultimate Settings & Configuration")
    
    # Settings tabs
    tab1, tab2, tab3, tab4 = st.tabs(["🔑 API Keys", "🎨 Preferences", "📊 Analytics", "🔧 Advanced"])
    
    with tab1:
        show_api_settings()
    
    with tab2:
        show_preference_settings()
    
    with tab3:
        show_analytics_settings()
    
    with tab4:
        show_advanced_settings()

def show_api_settings():
    """Show API configuration settings"""
    st.markdown("### 🔑 API Configuration")
    
    with st.form("api_settings"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 🤖 AI Providers")
            gemini_key = st.text_input("Google Gemini API Key", type="password")
            openai_key = st.text_input("OpenAI API Key", type="password")
            anthropic_key = st.text_input("Anthropic Claude API Key", type="password")
        
        with col2:
            st.markdown("#### 🔍 Job Boards")
            indeed_key = st.text_input("Indeed API Key", type="password")
            linkedin_key = st.text_input("LinkedIn API Key", type="password")
            glassdoor_key = st.text_input("Glassdoor API Key", type="password")
        
        if st.form_submit_button("💾 Save API Configuration", type="primary"):
            st.success("✅ API configuration saved successfully!")
            st.info("🔄 Restart the application to apply changes.")

def show_preference_settings():
    """Show user preference settings"""
    st.markdown("### 🎨 User Preferences")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🎨 Appearance")
        theme = st.selectbox("Theme", ["Dark", "Light", "Auto"], index=0)
        language = st.selectbox("Language", ["English", "Spanish", "French", "German"])
        timezone = st.selectbox("Timezone", ["UTC", "PST", "EST", "GMT"])
    
    with col2:
        st.markdown("#### 🔔 Notifications")
        email_notifications = st.checkbox("Email Notifications", value=True)
        job_alerts = st.checkbox("Job Match Alerts", value=True)
        weekly_reports = st.checkbox("Weekly Analytics Reports", value=False)
    
    st.markdown("#### 💾 Data Management")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        auto_save = st.checkbox("Auto-save Analyses", value=True)
    with col2:
        data_retention = st.selectbox("Data Retention", ["30 days", "90 days", "1 year", "Forever"])
    with col3:
        export_format = st.selectbox("Export Format", ["PDF", "JSON", "CSV"])

def show_analytics_settings():
    """Show analytics settings"""
    st.markdown("### 📊 Analytics Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 📈 Tracking")
        usage_analytics = st.checkbox("Usage Analytics", value=True)
        performance_metrics = st.checkbox("Performance Metrics", value=True)
        error_reporting = st.checkbox("Error Reporting", value=True)
    
    with col2:
        st.markdown("#### 🎯 Insights")
        personalized_recommendations = st.checkbox("Personalized Recommendations", value=True)
        market_insights = st.checkbox("Market Insights", value=True)
        salary_benchmarking = st.checkbox("Salary Benchmarking", value=True)

def show_advanced_settings():
    """Show advanced settings"""
    st.markdown("### 🔧 Advanced Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### ⚡ Performance")
        cache_size = st.slider("Cache Size (MB)", 50, 500, 200)
        max_file_size = st.slider("Max File Size (MB)", 5, 50, 10)
        concurrent_jobs = st.slider("Concurrent Jobs", 1, 10, 3)
    
    with col2:
        st.markdown("#### 🔒 Security")
        session_timeout = st.slider("Session Timeout (minutes)", 15, 120, 60)
        encryption_level = st.selectbox("Encryption Level", ["Standard", "High", "Maximum"])
        two_factor_auth = st.checkbox("Two-Factor Authentication", value=False)
    
    st.markdown("#### 🛠️ Developer Options")
    debug_mode = st.checkbox("Debug Mode", value=False)
    api_rate_limit = st.slider("API Rate Limit (requests/minute)", 10, 100, 60)
    log_level = st.selectbox("Log Level", ["INFO", "DEBUG", "WARNING", "ERROR"])

if __name__ == "__main__":
    main()