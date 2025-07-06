"""
Enhanced File Reader Utility
============================

Robust file reading with support for multiple formats and fallback mechanisms.
Handles PDF, DOCX, TXT files with comprehensive error handling.
"""

import os
import logging
import tempfile
import io
from typing import Optional, Union

# Configure logging
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path_or_bytes: Union[str, bytes]) -> str:
    """
    Extract text from PDF file with multiple fallback methods.
    
    Args:
        file_path_or_bytes: File path string or bytes content
        
    Returns:
        str: Extracted text or error message
    """
    try:
        # Method 1: PyPDF2
        try:
            from PyPDF2 import PdfReader
            
            if isinstance(file_path_or_bytes, str):
                if not os.path.exists(file_path_or_bytes):
                    raise FileNotFoundError(f"PDF file not found at {file_path_or_bytes}")
                reader = PdfReader(file_path_or_bytes)
            else:
                reader = PdfReader(io.BytesIO(file_path_or_bytes))
            
            if len(reader.pages) == 0:
                return "The PDF file appears to be empty."
            
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                return text.strip()
            else:
                raise Exception("No text extracted with PyPDF2")
                
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")
            
            # Method 2: pdfplumber (if available)
            try:
                import pdfplumber
                
                if isinstance(file_path_or_bytes, str):
                    with pdfplumber.open(file_path_or_bytes) as pdf:
                        text = ""
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                else:
                    with pdfplumber.open(io.BytesIO(file_path_or_bytes)) as pdf:
                        text = ""
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                
                if text.strip():
                    return text.strip()
                else:
                    raise Exception("No text extracted with pdfplumber")
                    
            except ImportError:
                logger.warning("pdfplumber not available")
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
                
                # Method 3: pymupdf/fitz (if available)
                try:
                    import fitz  # PyMuPDF
                    
                    if isinstance(file_path_or_bytes, str):
                        doc = fitz.open(file_path_or_bytes)
                    else:
                        doc = fitz.open(stream=file_path_or_bytes, filetype="pdf")
                    
                    text = ""
                    for page in doc:
                        text += page.get_text() + "\n"
                    doc.close()
                    
                    if text.strip():
                        return text.strip()
                    else:
                        raise Exception("No text extracted with PyMuPDF")
                        
                except ImportError:
                    logger.warning("PyMuPDF not available")
                except Exception as e:
                    logger.warning(f"PyMuPDF failed: {e}")
                    
                    return "No text could be extracted from the PDF. It may be scanned or contain only images."
    
    except Exception as e:
        logger.error(f"All PDF extraction methods failed: {e}")
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file_path_or_bytes: Union[str, bytes]) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path_or_bytes: File path string or bytes content
        
    Returns:
        str: Extracted text or error message
    """
    try:
        from docx import Document
        
        if isinstance(file_path_or_bytes, str):
            if not os.path.exists(file_path_or_bytes):
                raise FileNotFoundError(f"DOCX file not found at {file_path_or_bytes}")
            doc = Document(file_path_or_bytes)
        else:
            doc = Document(io.BytesIO(file_path_or_bytes))
        
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if text.strip():
            return text.strip()
        else:
            return "The DOCX file appears to be empty or contains no readable text."
            
    except ImportError:
        return "python-docx library not installed. Cannot read DOCX files."
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return f"Error extracting text from DOCX: {str(e)}"

def extract_text_from_txt(file_path_or_bytes: Union[str, bytes]) -> str:
    """
    Extract text from TXT file with encoding detection.
    
    Args:
        file_path_or_bytes: File path string or bytes content
        
    Returns:
        str: Extracted text or error message
    """
    try:
        if isinstance(file_path_or_bytes, str):
            if not os.path.exists(file_path_or_bytes):
                raise FileNotFoundError(f"TXT file not found at {file_path_or_bytes}")
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path_or_bytes, 'r', encoding=encoding) as file:
                        text = file.read()
                        if text.strip():
                            return text.strip()
                except UnicodeDecodeError:
                    continue
            
            return "Could not decode the text file with any supported encoding."
        
        else:
            # Handle bytes content
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_path_or_bytes.decode(encoding)
                    if text.strip():
                        return text.strip()
                except UnicodeDecodeError:
                    continue
            
            return "Could not decode the text content with any supported encoding."
            
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        return f"Error extracting text from TXT: {str(e)}"

def extract_text_from_file(file_path_or_uploaded_file) -> str:
    """
    Universal file text extractor that handles multiple formats.
    
    Args:
        file_path_or_uploaded_file: File path string or Streamlit uploaded file object
        
    Returns:
        str: Extracted text or error message
    """
    try:
        # Handle Streamlit uploaded file
        if hasattr(file_path_or_uploaded_file, 'name') and hasattr(file_path_or_uploaded_file, 'getvalue'):
            file_name = file_path_or_uploaded_file.name
            file_content = file_path_or_uploaded_file.getvalue()
            file_extension = file_name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                return extract_text_from_pdf(file_content)
            elif file_extension == 'docx':
                return extract_text_from_docx(file_content)
            elif file_extension in ['txt', 'text']:
                return extract_text_from_txt(file_content)
            else:
                return f"Unsupported file format: {file_extension}. Supported formats: PDF, DOCX, TXT"
        
        # Handle file path
        elif isinstance(file_path_or_uploaded_file, str):
            if not os.path.exists(file_path_or_uploaded_file):
                return f"File not found: {file_path_or_uploaded_file}"
            
            file_extension = file_path_or_uploaded_file.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                return extract_text_from_pdf(file_path_or_uploaded_file)
            elif file_extension == 'docx':
                return extract_text_from_docx(file_path_or_uploaded_file)
            elif file_extension in ['txt', 'text']:
                return extract_text_from_txt(file_path_or_uploaded_file)
            else:
                return f"Unsupported file format: {file_extension}. Supported formats: PDF, DOCX, TXT"
        
        else:
            return "Invalid file input. Please provide a file path or uploaded file object."
            
    except Exception as e:
        logger.error(f"Error in universal file extractor: {e}")
        return f"Error processing file: {str(e)}"

def validate_file_size(file_size_bytes: int, max_size_mb: int = 10) -> bool:
    """
    Validate file size.
    
    Args:
        file_size_bytes: File size in bytes
        max_size_mb: Maximum allowed size in MB
        
    Returns:
        bool: True if file size is acceptable
    """
    max_size_bytes = max_size_mb * 1024 * 1024
    return file_size_bytes <= max_size_bytes

def get_file_info(file_path_or_uploaded_file) -> dict:
    """
    Get file information including size, type, and name.
    
    Args:
        file_path_or_uploaded_file: File path string or Streamlit uploaded file object
        
    Returns:
        dict: File information
    """
    try:
        if hasattr(file_path_or_uploaded_file, 'name'):
            # Streamlit uploaded file
            return {
                'name': file_path_or_uploaded_file.name,
                'size': len(file_path_or_uploaded_file.getvalue()),
                'type': file_path_or_uploaded_file.name.split('.')[-1].lower(),
                'size_mb': round(len(file_path_or_uploaded_file.getvalue()) / (1024 * 1024), 2)
            }
        elif isinstance(file_path_or_uploaded_file, str):
            # File path
            if os.path.exists(file_path_or_uploaded_file):
                size = os.path.getsize(file_path_or_uploaded_file)
                return {
                    'name': os.path.basename(file_path_or_uploaded_file),
                    'size': size,
                    'type': file_path_or_uploaded_file.split('.')[-1].lower(),
                    'size_mb': round(size / (1024 * 1024), 2)
                }
            else:
                return {'error': 'File not found'}
        else:
            return {'error': 'Invalid file input'}
            
    except Exception as e:
        return {'error': f'Error getting file info: {str(e)}'}

# Sample resume text for testing
SAMPLE_RESUME_TEXT = """
John Smith
Senior Software Engineer

Contact Information:
Email: john.smith@email.com
Phone: (555) 123-4567
LinkedIn: linkedin.com/in/johnsmith
Location: San Francisco, CA

PROFESSIONAL SUMMARY:
Experienced Senior Software Engineer with 8+ years of expertise in full-stack development, 
cloud architecture, and team leadership. Proven track record of delivering scalable 
solutions and mentoring development teams.

TECHNICAL SKILLS:
Programming Languages: Python, JavaScript, Java, TypeScript, Go
Web Technologies: React, Node.js, Express, HTML5, CSS3, REST APIs
Databases: PostgreSQL, MongoDB, Redis, MySQL
Cloud Platforms: AWS, Azure, Google Cloud Platform
DevOps: Docker, Kubernetes, Jenkins, CI/CD, Terraform
Tools: Git, JIRA, Confluence, VS Code, IntelliJ

PROFESSIONAL EXPERIENCE:

Senior Software Engineer | TechCorp Inc. | 2020 - Present
• Led development of microservices architecture serving 1M+ users
• Implemented machine learning models for recommendation systems
• Managed team of 5 developers and improved deployment frequency by 300%
• Reduced system latency by 40% through performance optimization
• Technologies: Python, React, AWS, Docker, Kubernetes

Software Engineer | StartupXYZ | 2018 - 2020
• Developed full-stack web applications using React and Node.js
• Built RESTful APIs handling 10k+ requests per minute
• Implemented automated testing reducing bugs by 60%
• Collaborated with product team on feature requirements
• Technologies: JavaScript, React, Node.js, MongoDB, AWS

Junior Developer | WebSolutions | 2016 - 2018
• Developed responsive websites and web applications
• Worked with SQL databases and implemented data visualization
• Participated in agile development processes
• Technologies: HTML, CSS, JavaScript, PHP, MySQL

EDUCATION:
Master of Science in Computer Science
Stanford University | 2014 - 2016
GPA: 3.8/4.0

Bachelor of Science in Software Engineering
University of California, Berkeley | 2010 - 2014
GPA: 3.7/4.0

CERTIFICATIONS:
• AWS Certified Solutions Architect - Professional (2023)
• Google Cloud Professional Developer (2022)
• Certified Kubernetes Administrator (2021)
• Scrum Master Certification (2020)

PROJECTS:
E-commerce Platform (2023)
• Built scalable e-commerce platform handling 50k+ daily transactions
• Implemented real-time inventory management and payment processing
• Technologies: Python, React, PostgreSQL, Redis, AWS

Data Analytics Dashboard (2022)
• Created real-time analytics dashboard for business intelligence
• Integrated multiple data sources and implemented data pipelines
• Technologies: Python, D3.js, Apache Kafka, Elasticsearch

ACHIEVEMENTS:
• Employee of the Year 2022 at TechCorp Inc.
• Led team that won company hackathon 2021
• Speaker at PyCon 2023 on "Scalable Microservices"
• Contributed to 5+ open-source projects with 1000+ GitHub stars

LANGUAGES:
• English (Native)
• Spanish (Conversational)
• Mandarin (Basic)
"""